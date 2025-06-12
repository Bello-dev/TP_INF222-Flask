'''
Script pour générer un rapport DOCX structuré à partir du contenu de README.md.
'''
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def extract_section_content(readme_lines, section_title_text):
    """
    Extrait le contenu d'une section spécifique du README.
    La section commence par un titre (ex: "## Titre de la Section")
    et se termine avant le prochain titre de niveau égal ou supérieur, ou la fin du fichier.
    """
    content_lines = []
    in_section = False
    section_heading_level = 0
    start_index = -1

    # Trouver la section cible et son niveau
    for i, line in enumerate(readme_lines):
        stripped_line = line.strip()
        match = re.match(r"^(#{1,6})\s+(.*)", stripped_line)
        if match:
            level = len(match.group(1))
            title = match.group(2).strip()
            if title == section_title_text:
                in_section = True
                section_heading_level = level
                start_index = i + 1  # Le contenu commence après la ligne du titre
                break
    
    if not in_section or start_index == -1:
        return ""

    # Extraire le contenu jusqu'au prochain titre de niveau égal/supérieur ou EOF
    for i in range(start_index, len(readme_lines)):
        line = readme_lines[i]
        stripped_line = line.strip()
        
        next_heading_match = re.match(r"^(#{1,6})\s+.*", stripped_line)
        if next_heading_match:
            next_heading_level = len(next_heading_match.group(1))
            if next_heading_level <= section_heading_level:
                break  # Fin de la section actuelle
        
        content_lines.append(line) 
        
    return "".join(content_lines).strip()

def add_content_to_doc(doc, text_content, is_list_candidate=False):
    """Ajoute le contenu textuel au document, gérant les paragraphes, listes et blocs de code."""
    if not text_content:
        return

    lines = text_content.split('\n')
    is_code_block = False
    code_block_lines = []

    for line_raw in lines:
        line = line_raw.strip()

        if line.startswith("```"):
            if not is_code_block:
                is_code_block = True
            else:
                is_code_block = False
                if code_block_lines:
                    # S'assure que le paragraphe est ajouté avant le run pour éviter les erreurs
                    p = doc.add_paragraph()
                    run = p.add_run("\n".join(code_block_lines))
                    run.font.name = 'Courier New'
                    code_block_lines = []
            continue

        if is_code_block:
            code_block_lines.append(line_raw) # Conserver l'indentation pour le code
            continue

        if not line.strip(): # Ligne vide
            # doc.add_paragraph() # Ajoute un espacement, peut être géré par les styles
            continue


        # Détection basique de listes
        if line.startswith(("* ", "- ")) and is_list_candidate:
            item_text = line[2:].strip()
            doc.add_paragraph(item_text, style='ListBullet')
        elif re.match(r"^\d+\.\s+", line) and is_list_candidate:
            item_text = re.sub(r"^\d+\.\s+", "", line).strip()
            doc.add_paragraph(item_text, style='ListNumber')
        else:
            doc.add_paragraph(line) # Paragraphe normal

    if is_code_block and code_block_lines: # Bloc de code non terminé à la fin
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_block_lines))
        run.font.name = 'Courier New'

# --- Lecture du README.md ---
try:
    with open('README.md', 'r', encoding='utf-8') as f:
        readme_lines = f.readlines()
except FileNotFoundError:
    print("Erreur : README.md non trouvé.")
    exit(1)

doc = Document()

# --- Page de Titre ---
project_title_text = "Rapport de Projet : TP_INF222-Flask" 
# Essayer d'extraire le titre H1 du README
h1_match = next((line for line in readme_lines if line.strip().startswith("# ")), None)
if h1_match:
    project_title_text = "Rapport de Projet : " + h1_match.strip("# ").strip()

title_heading = doc.add_heading(project_title_text, level=0)
title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

date_paragraph = doc.add_paragraph()
date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_paragraph.add_run(f"Date de génération : {datetime.now().strftime('%d %B %Y')}").italic = True

auteurs_content_raw = extract_section_content(readme_lines, "Auteurs")
if auteurs_content_raw:
    auteurs_paragraph = doc.add_paragraph()
    auteurs_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Nettoyer un peu le contenu des auteurs (supprimer les puces si présentes)
    cleaned_auteurs = "\n".join([l.lstrip("-* ").strip() for l in auteurs_content_raw.split('\n') if l.strip()])
    auteurs_paragraph.add_run("Auteur(s) : " + cleaned_auteurs.replace("\n", ", ")).italic = True

doc.add_page_break()

# --- Sections du Rapport ---
# (Titre dans Word, Titre exact dans README après "## ", Est-ce une liste potentielle?)
sections_map = [
    {"doc_title": "Introduction", "readme_title": "Description", "is_list": False},
    {"doc_title": "Fonctionnalités", "readme_title": "Fonctionnalités", "is_list": True},
    {"doc_title": "Architecture du Projet", "readme_title": "Architecture", "is_list": False},
    {"doc_title": "Guide d'Installation", "readme_title": "Installation", "is_list": True}, # Installation a souvent des étapes listées
    {"doc_title": "Guide d'Utilisation", "readme_title": "Utilisation (Endpoints API)", "is_list": False},
    {"doc_title": "Stratégie de Tests", "readme_title": "Tests", "is_list": False},
    {"doc_title": "Déploiement", "readme_title": "Déploiement", "is_list": False},
    {"doc_title": "Licence", "readme_title": "Licence", "is_list": False}
]

section_counter = 1
for section_info in sections_map:
    doc_section_title = section_info["doc_title"]
    readme_section_title_text = section_info["readme_title"]
    is_list_candidate = section_info["is_list"]

    content = extract_section_content(readme_lines, readme_section_title_text)
    
    if content or doc_section_title == "Introduction": # Toujours ajouter l'intro
        doc.add_heading(f"{section_counter}. {doc_section_title}", level=1)
        if not content and doc_section_title == "Introduction":
            doc.add_paragraph(f"Ce document présente une vue d'ensemble du projet {project_title_text.replace('Rapport de Projet : ', '')}.")
        
        add_content_to_doc(doc, content, is_list_candidate=is_list_candidate)
        # Ajouter un espacement après la section si du contenu a été ajouté
        if content:
             doc.add_paragraph()
        section_counter += 1

# --- Sauvegarde du Document ---
output_filename = "TP_INF222-Flask_Rapport_Structure.docx"
try:
    doc.save(output_filename)
    print(f"Rapport structuré généré avec succès : {output_filename}")
except Exception as e:
    print(f"Erreur lors de la sauvegarde du document : {e}")
    exit(1)
